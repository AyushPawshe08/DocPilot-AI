import logging
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.exceptions import DocumentGenerationError, DocumentNotFoundError

logger = logging.getLogger(__name__)

_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


class DocumentService:

    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(self, title: str, content: str) -> tuple[str, str]:
        try:
            document = Document()

            heading = document.add_heading(title, level=0)
            if heading.runs:
                heading.runs[0].font.size = Pt(24)

            for raw_line in content.split("\n"):
                line = raw_line.strip()
                if not line:
                    continue
                self._add_line(document, line)

            filename = f"{self._safe_filename(title)}.docx"
            path = self.output_dir / filename

            # Prevent overwriting an existing file
            counter = 1
            while path.exists():
                filename = f"{self._safe_filename(title)} ({counter}).docx"
                path = self.output_dir / filename
                counter += 1

            document.save(path)

        except Exception as exc:
            logger.exception("Document rendering failed")
            raise DocumentGenerationError(
                f"Failed to render .docx file: {exc}"
            ) from exc

        logger.info("Document rendered: %s", filename)
        return filename, str(path)

    def resolve_path(self, filename: str) -> Path:
        filename = Path(filename).name
        path = self.output_dir / filename

        if not path.exists():
            raise DocumentNotFoundError(
                f"No document found: {filename}"
            )

        return path

    @staticmethod
    def _safe_filename(title: str) -> str:
        title = re.sub(r'[<>:"/\\|?*]', "", title)
        title = re.sub(r"\s+", " ", title).strip()
        return title[:100]

    @staticmethod
    def _add_line(document: Document, line: str) -> None:
        if line.startswith("## "):
            document.add_heading(line[3:], level=2)

        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)

        elif line.startswith(("- ", "* ")):
            DocumentService._add_runs(
                document.add_paragraph(style="List Bullet"),
                line[2:],
            )

        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            DocumentService._add_runs(
                document.add_paragraph(style="List Number"),
                text,
            )

        else:
            DocumentService._add_runs(
                document.add_paragraph(),
                line,
            )

    @staticmethod
    def _add_runs(paragraph, text: str) -> None:
        pos = 0

        for match in _BOLD_PATTERN.finditer(text):
            if match.start() > pos:
                paragraph.add_run(text[pos:match.start()])

            paragraph.add_run(match.group(1)).bold = True
            pos = match.end()

        if pos < len(text):
            paragraph.add_run(text[pos:])