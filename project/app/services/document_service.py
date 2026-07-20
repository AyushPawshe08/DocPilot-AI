"""Renders Markdown-flavoured agent output into a .docx file.

Kept deliberately dependency-light (python-docx only) and deterministic:
this stage never calls the LLM, so it is fully unit-testable and cannot
itself introduce non-determinism or extra latency/cost into the pipeline.
"""
import logging
import re
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.shared import Pt

from app.exceptions import DocumentGenerationError, DocumentNotFoundError

logger = logging.getLogger(__name__)

_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


class DocumentService:
    """Converts (title, markdown content) into a stored .docx file.

    Files are named ``<document_id>.docx`` and ``document_id`` is a UUID
    generated here -- the API never trusts a client-supplied filename,
    which also closes the path-traversal hole present in naive
    "download by filename" implementations.
    """

    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(self, title: str, content: str) -> tuple[str, str]:
        """Render the document and return ``(document_id, file_path)``."""
        try:
            document = Document()

            heading = document.add_heading(title, level=0)
            if heading.runs:
                heading.runs[0].font.size = Pt(24)

            for raw_line in content.split("\n"):
                line = raw_line.strip()
                if not line:
                    continue
                self._add_line(document, line)

            document_id = uuid4().hex
            path = self.output_dir / f"{document_id}.docx"
            document.save(path)
        except Exception as exc:  # noqa: BLE001
            logger.exception("Document rendering failed")
            raise DocumentGenerationError(f"Failed to render .docx file: {exc}") from exc

        logger.info("Document rendered: id=%s path=%s", document_id, path)
        return document_id, str(path)

    def resolve_path(self, document_id: str) -> Path:
        """Return the file path for a previously generated document.

        Validates the id is a plain hex UUID before touching the
        filesystem so this can never be used to read arbitrary files.
        """
        if not re.fullmatch(r"[0-9a-f]{32}", document_id):
            raise DocumentNotFoundError(f"Invalid document id: {document_id!r}")

        path = self.output_dir / f"{document_id}.docx"
        if not path.is_file():
            raise DocumentNotFoundError(f"No document found for id: {document_id!r}")
        return path

    @staticmethod
    def _add_line(document: Document, line: str) -> None:
        if line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            DocumentService._add_runs(
                document.add_paragraph(style="List Bullet"), line[2:]
            )
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            DocumentService._add_runs(
                document.add_paragraph(style="List Number"), text
            )
        else:
            DocumentService._add_runs(document.add_paragraph(), line)

    @staticmethod
    def _add_runs(paragraph, text: str) -> None:
        """Split ``**bold**`` markdown into bold/non-bold runs."""
        pos = 0
        for match in _BOLD_PATTERN.finditer(text):
            if match.start() > pos:
                paragraph.add_run(text[pos:match.start()])
            paragraph.add_run(match.group(1)).bold = True
            pos = match.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])
